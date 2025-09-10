"""
문서 파싱 서비스
PDF, DOCX, TXT 파일에서 텍스트를 추출하는 서비스
"""
import os
import mimetypes
from typing import Dict, Any, Optional, Tuple
from pathlib import Path
import logging

# 문서 파싱 라이브러리
import PyPDF2
import pdfplumber
from docx import Document
import tiktoken

logger = logging.getLogger(__name__)

class DocumentParser:
    """문서 파싱 클래스"""
    
    def __init__(self):
        # tiktoken 인코더 초기화 (토큰 계산용)
        self.tokenizer = tiktoken.get_encoding("cl100k_base")
    
    def parse_document(self, file_path: str, content_type: str) -> Dict[str, Any]:
        """
        문서를 파싱하여 텍스트와 메타데이터를 추출
        
        Args:
            file_path: 파일 경로
            content_type: MIME 타입
            
        Returns:
            Dict containing:
                - text: 추출된 텍스트
                - metadata: 문서 메타데이터
                - token_count: 토큰 수
                - page_count: 페이지 수 (PDF인 경우)
        """
        try:
            if content_type == "application/pdf":
                return self._parse_pdf(file_path)
            elif content_type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
                return self._parse_docx(file_path)
            elif content_type == "text/plain":
                return self._parse_txt(file_path)
            else:
                raise ValueError(f"지원하지 않는 파일 형식: {content_type}")
                
        except Exception as e:
            logger.error(f"문서 파싱 실패: {file_path}, 에러: {str(e)}")
            raise
    
    def _parse_pdf(self, file_path: str) -> Dict[str, Any]:
        """PDF 파일 파싱"""
        try:
            # pdfplumber를 사용한 고품질 텍스트 추출
            with pdfplumber.open(file_path) as pdf:
                text_parts = []
                page_count = len(pdf.pages)
                
                for page_num, page in enumerate(pdf.pages, 1):
                    page_text = page.extract_text()
                    if page_text:
                        text_parts.append(f"[페이지 {page_num}]\n{page_text}")
                
                full_text = "\n\n".join(text_parts)
                
                # 메타데이터 추출
                metadata = {
                    "page_count": page_count,
                    "title": pdf.metadata.get("Title", ""),
                    "author": pdf.metadata.get("Author", ""),
                    "subject": pdf.metadata.get("Subject", ""),
                    "creator": pdf.metadata.get("Creator", ""),
                    "producer": pdf.metadata.get("Producer", ""),
                    "creation_date": str(pdf.metadata.get("CreationDate", "")),
                    "modification_date": str(pdf.metadata.get("ModDate", ""))
                }
                
                # 토큰 수 계산
                token_count = len(self.tokenizer.encode(full_text))
                
                return {
                    "text": full_text,
                    "metadata": metadata,
                    "token_count": token_count,
                    "page_count": page_count
                }
                
        except Exception as e:
            logger.error(f"PDF 파싱 실패: {file_path}, 에러: {str(e)}")
            # PyPDF2로 폴백 시도
            return self._parse_pdf_fallback(file_path)
    
    def _parse_pdf_fallback(self, file_path: str) -> Dict[str, Any]:
        """PyPDF2를 사용한 PDF 파싱 (폴백)"""
        try:
            with open(file_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                text_parts = []
                page_count = len(pdf_reader.pages)
                
                for page_num, page in enumerate(pdf_reader.pages, 1):
                    page_text = page.extract_text()
                    if page_text:
                        text_parts.append(f"[페이지 {page_num}]\n{page_text}")
                
                full_text = "\n\n".join(text_parts)
                
                # 메타데이터 추출
                metadata = {
                    "page_count": page_count,
                    "title": pdf_reader.metadata.get("/Title", "") if pdf_reader.metadata else "",
                    "author": pdf_reader.metadata.get("/Author", "") if pdf_reader.metadata else "",
                    "subject": pdf_reader.metadata.get("/Subject", "") if pdf_reader.metadata else "",
                    "creator": pdf_reader.metadata.get("/Creator", "") if pdf_reader.metadata else "",
                    "producer": pdf_reader.metadata.get("/Producer", "") if pdf_reader.metadata else "",
                    "creation_date": str(pdf_reader.metadata.get("/CreationDate", "")) if pdf_reader.metadata else "",
                    "modification_date": str(pdf_reader.metadata.get("/ModDate", "")) if pdf_reader.metadata else ""
                }
                
                token_count = len(self.tokenizer.encode(full_text))
                
                return {
                    "text": full_text,
                    "metadata": metadata,
                    "token_count": token_count,
                    "page_count": page_count
                }
                
        except Exception as e:
            logger.error(f"PDF 폴백 파싱도 실패: {file_path}, 에러: {str(e)}")
            raise
    
    def _parse_docx(self, file_path: str) -> Dict[str, Any]:
        """DOCX 파일 파싱"""
        try:
            doc = Document(file_path)
            
            # 텍스트 추출
            text_parts = []
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text_parts.append(paragraph.text)
            
            # 표에서 텍스트 추출
            for table in doc.tables:
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        if cell.text.strip():
                            row_text.append(cell.text.strip())
                    if row_text:
                        text_parts.append(" | ".join(row_text))
            
            full_text = "\n".join(text_parts)
            
            # 메타데이터 추출
            core_props = doc.core_properties
            metadata = {
                "title": core_props.title or "",
                "author": core_props.author or "",
                "subject": core_props.subject or "",
                "keywords": core_props.keywords or "",
                "comments": core_props.comments or "",
                "created": str(core_props.created) if core_props.created else "",
                "modified": str(core_props.modified) if core_props.modified else "",
                "last_modified_by": core_props.last_modified_by or ""
            }
            
            token_count = len(self.tokenizer.encode(full_text))
            
            return {
                "text": full_text,
                "metadata": metadata,
                "token_count": token_count,
                "page_count": 1  # DOCX는 페이지 수를 정확히 알기 어려움
            }
            
        except Exception as e:
            logger.error(f"DOCX 파싱 실패: {file_path}, 에러: {str(e)}")
            raise
    
    def _parse_txt(self, file_path: str) -> Dict[str, Any]:
        """TXT 파일 파싱"""
        try:
            # 다양한 인코딩 시도
            encodings = ['utf-8', 'cp949', 'euc-kr', 'latin-1']
            text = None
            
            for encoding in encodings:
                try:
                    with open(file_path, 'r', encoding=encoding) as file:
                        text = file.read()
                        break
                except UnicodeDecodeError:
                    continue
            
            if text is None:
                raise ValueError("파일 인코딩을 감지할 수 없습니다.")
            
            # 메타데이터
            file_stat = os.stat(file_path)
            metadata = {
                "file_size": file_stat.st_size,
                "encoding": encoding,
                "created": str(file_stat.st_ctime),
                "modified": str(file_stat.st_mtime)
            }
            
            token_count = len(self.tokenizer.encode(text))
            
            return {
                "text": text,
                "metadata": metadata,
                "token_count": token_count,
                "page_count": 1
            }
            
        except Exception as e:
            logger.error(f"TXT 파싱 실패: {file_path}, 에러: {str(e)}")
            raise
    
    def get_supported_formats(self) -> Dict[str, str]:
        """지원하는 파일 형식 반환"""
        return {
            "application/pdf": "PDF 문서",
            "application/vnd.openxmlformats-officedocument.wordprocessingml.document": "Word 문서 (DOCX)",
            "text/plain": "텍스트 파일 (TXT)"
        }
    
    def validate_file(self, file_path: str, content_type: str) -> bool:
        """파일 유효성 검증"""
        try:
            if not os.path.exists(file_path):
                return False
            
            # 파일 크기 검증 (100MB 제한)
            file_size = os.path.getsize(file_path)
            if file_size > 100 * 1024 * 1024:  # 100MB
                return False
            
            # MIME 타입 검증
            supported_formats = self.get_supported_formats()
            if content_type not in supported_formats:
                return False
            
            return True
            
        except Exception as e:
            logger.error(f"파일 검증 실패: {file_path}, 에러: {str(e)}")
            return False
